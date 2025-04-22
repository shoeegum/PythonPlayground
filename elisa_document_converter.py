import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement, qn
from docx.table import Table
# The following import is not needed as we'll use a more compatible way to set line spacing

class ELISAProcessor:
    def __init__(self, template_path, supplier_path):
        self.template = Document(template_path)
        self.supplier = Document(supplier_path)
        self.supplier_filename = os.path.splitext(os.path.basename(supplier_path))[0]  # get 'supplier1' from 'supplier1.docx'
        self.company = self._detect_company(supplier_path)
        self.content = {'sections': {}, 'tables': [], 'images': []}
        self.mappings = self._load_mappings()
        self._styles = {
            'header': {'font': 'Calibri', 'size': 24, 'bold': True, 
                      'color': RGBColor(0x00, 0x00, 0xFF), 'alignment': WD_ALIGN_PARAGRAPH.RIGHT},
            'body': {'font': 'Calibri', 'size': 11, 
                     'line_spacing': 1.15, 'color': RGBColor(0x00, 0x00, 0x00)},
            'disclaimer': {'italic': True}
        }
        # Initialize with default replacements
        self._replacements = []
        # Add default Boster replacement if needed
        if self.company == 'boster':
            self._replacements.append(('Boster', 'Innovative Research'))

    def _detect_company(self, path):
        filename = os.path.basename(path)
        # Simple filename-based detection
        if filename.startswith('EK'): return 'boster'
        if filename.startswith('RDR'): return 'red_dot'
        if filename.startswith('AP'): return 'assay_pro'
        if filename.startswith('ICL'): return 'icl'
        if filename.startswith('EL'): return 'elabs'
        if filename.startswith('AA'): return 'arbor_assay'
        
        # Fallback to content-based detection
        try:
            with open(path, 'rb') as f:
                content = f.read().decode('utf-8', errors='ignore')
                if "Boster" in content: return 'boster'
                if "Red Dot" in content: return 'red_dot'
                if "Assay Pro" in content: return 'assay_pro'
                if "ICL" in content: return 'icl'
                if "eLabs" in content: return 'elabs'
                if "Arbor Assay" in content: return 'arbor_assay'
        except Exception:
            pass
            
        # Use default if detection fails
        return list(self._load_mappings().keys())[0]

    def _load_mappings(self):
        return {
            'boster': {
                'section_map': {
                    'INTENDED USE': {'source': r'Assay Principle', 'regex': True, 'paragraph': 0},
                    'BACKGROUND': {'source': r'Background on', 'regex': True},
                    'ASSAY PRINCIPLE': {'source': r'Assay Principle', 'regex': True, 'paragraph': -1},
                    'OVERVIEW': {'source': 'Overview', 'table_filter': ['Product Name', 'Description']},
                    'TECHNICAL DETAILS': {'source': 'Technical Details'},
                    'STANDARD CURVE EXAMPLE': {'source': r'Standard Curve Example$', 'regex': True},
                    'INTRA/INTER-ASSAY VARIABILITY': {
                        'source': r'(Intra|Inter).*Variability',
                        'regex': True,
                        'table_headers': ['Intra-Assay', 'Inter-Assay']
                    },
                    'REPRODUCIBILITY': {
                        'source': r'Reproducibility',
                        'table_headers': ['Reproducibility']
                    },
                    'REAGENT PREPARATION AND STORAGE': {
                        'source': 'Reagent Preparation',
                        'special_handling': True,
                        'table_headers': ['Component', 'Preparation', 'Storage']
                    }
                },
                'remove_sections': ['Dilution of Standard', 'Sample Collection Notes', 'Sample Activation']
            },
            'red_dot': {
                'section_map': {
                    'INTENDED USE': {'source': r'Assay Principle', 'regex': True, 'paragraph': 0},
                    'BACKGROUND': {'source': r'Background on', 'regex': True},
                    'ASSAY PRINCIPLE': {'source': r'Assay Principle', 'regex': True, 'paragraph': -1},
                    'OVERVIEW': {'source': 'Overview', 'table_filter': ['Product Name', 'Description']},
                    'TECHNICAL DETAILS': {'source': 'Technical Details'},
                    'STANDARD CURVE EXAMPLE': {'source': r'Standard Curve Example$', 'regex': True},
                    'INTRA/INTER-ASSAY VARIABILITY': {
                        'source': r'(Intra|Inter).*Variability',
                        'regex': True,
                        'table_headers': ['Intra-Assay', 'Inter-Assay']
                    },
                    'REPRODUCIBILITY': {
                        'source': r'Reproducibility',
                        'table_headers': ['Reproducibility']
                    }
                },
                'remove_sections': ['Dilution of Standard', 'Sample Collection Notes', 'Sample Activation']
            },
            'assay_pro': {
                'section_map': {
                    'INTENDED USE': {'source': r'Intended Use|Purpose', 'regex': True, 'paragraph': 0},
                    'BACKGROUND': {'source': r'Background|Introduction', 'regex': True},
                    'ASSAY PRINCIPLE': {'source': r'Assay Principle|Principle of the Assay', 'regex': True},
                    'OVERVIEW': {'source': r'Overview|Kit Contents', 'regex': True},
                    'TECHNICAL DETAILS': {'source': r'Technical Details|Specifications', 'regex': True},
                    'STANDARD CURVE EXAMPLE': {'source': r'Standard Curve|Curve Example', 'regex': True},
                    'INTRA/INTER-ASSAY VARIABILITY': {
                        'source': r'(Intra|Inter).*Variability|Precision',
                        'regex': True,
                        'table_headers': ['Intra-Assay', 'Inter-Assay', 'Precision']
                    },
                    'REPRODUCIBILITY': {
                        'source': r'Reproducibility',
                        'table_headers': ['Reproducibility']
                    },
                    'REAGENT PREPARATION AND STORAGE': {
                        'source': r'Reagent Preparation|Kit Preparation',
                        'regex': True,
                        'special_handling': True,
                        'table_headers': ['Component', 'Preparation', 'Storage']
                    }
                },
                'remove_sections': ['Sample Collection Notes', 'Sample Activation']
            },
            'icl': {
                'section_map': {
                    'INTENDED USE': {'source': r'Intended Use|Kit Application', 'regex': True, 'paragraph': 0},
                    'BACKGROUND': {'source': r'Background|Introduction', 'regex': True},
                    'ASSAY PRINCIPLE': {'source': r'Assay Principle|Method', 'regex': True},
                    'OVERVIEW': {'source': r'Overview|Summary', 'regex': True},
                    'TECHNICAL DETAILS': {'source': r'Technical Details|Specifications', 'regex': True},
                    'STANDARD CURVE EXAMPLE': {'source': r'Standard Curve|Calibration', 'regex': True},
                    'INTRA/INTER-ASSAY VARIABILITY': {
                        'source': r'(Intra|Inter).*Variability|Precision|CV',
                        'regex': True,
                        'table_headers': ['Intra-Assay', 'Inter-Assay', 'CV']
                    },
                    'REAGENT PREPARATION AND STORAGE': {
                        'source': r'Reagent Preparation|Materials|Components',
                        'regex': True,
                        'special_handling': True,
                        'table_headers': ['Component', 'Preparation', 'Storage']
                    }
                },
                'remove_sections': ['Sample Dilution', 'Sample Activation']
            },
            'elabs': {
                'section_map': {
                    'INTENDED USE': {'source': r'Intended Use|Purpose', 'regex': True, 'paragraph': 0},
                    'BACKGROUND': {'source': r'Background|Introduction', 'regex': True},
                    'ASSAY PRINCIPLE': {'source': r'Assay Principle|Test Principle', 'regex': True},
                    'OVERVIEW': {'source': r'Overview|Kit Components', 'regex': True},
                    'TECHNICAL DETAILS': {'source': r'Technical Details|Specifications', 'regex': True},
                    'STANDARD CURVE EXAMPLE': {'source': r'Standard Curve|Calibration Curve', 'regex': True},
                    'INTRA/INTER-ASSAY VARIABILITY': {
                        'source': r'(Intra|Inter).*Variability|Precision|Variation',
                        'regex': True,
                        'table_headers': ['Intra-Assay', 'Inter-Assay', 'Variation']
                    },
                    'REAGENT PREPARATION AND STORAGE': {
                        'source': r'Reagent Preparation|Reagents|Materials',
                        'regex': True,
                        'special_handling': True,
                        'table_headers': ['Component', 'Preparation', 'Storage']
                    }
                },
                'remove_sections': ['Sample Collection Notes']
            },
            'arbor_assay': {
                'section_map': {
                    'INTENDED USE': {'source': r'Intended Use|Description', 'regex': True, 'paragraph': 0},
                    'BACKGROUND': {'source': r'Background|Introduction', 'regex': True},
                    'ASSAY PRINCIPLE': {'source': r'Assay Principle|Principle', 'regex': True},
                    'OVERVIEW': {'source': r'Overview|Contents|Kit Contents', 'regex': True},
                    'TECHNICAL DETAILS': {'source': r'Technical Details|Performance|Specifications', 'regex': True},
                    'STANDARD CURVE EXAMPLE': {'source': r'Standard Curve|Typical Data', 'regex': True},
                    'INTRA/INTER-ASSAY VARIABILITY': {
                        'source': r'(Intra|Inter).*Variability|Precision|Performance',
                        'regex': True,
                        'table_headers': ['Intra-Assay', 'Inter-Assay', 'Performance']
                    },
                    'REAGENT PREPARATION AND STORAGE': {
                        'source': r'Reagent Preparation|Materials|Kit Contents',
                        'regex': True,
                        'special_handling': True,
                        'table_headers': ['Component', 'Preparation', 'Storage']
                    }
                },
                'remove_sections': ['Sample Activation']
            },
            'default': {'section_map': {}, 'remove_sections': []}
        }

    def _extract_sections(self):
        sections = {}
        current_section = None
        current_content = []

        for para in self.supplier.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith('Heading') or self._is_likely_heading(para):
                if current_section is not None:
                    sections[current_section] = current_content
                current_section = text
                current_content = [para]
            else:
                if current_section is not None:
                    current_content.append(para)

        if current_section is not None:
            sections[current_section] = current_content
        self.content['sections'] = sections
        
        for table in self.supplier.tables:
            table_section = self._find_table_section(table)
            if table_section:
                if table_section not in self.content.get('tables',{}):
                    self.content['tables'][table_section] = []
                self.content['tables'][table_section].append(table)


    def _is_likely_heading(self, para):
        if para.runs and para.runs[0].bold:
            return True
        if len(para.text) < 100 and para.text.strip().endswith(':'):
            return True
        heading_patterns = [
            r'^[A-Z][a-zA-Z\s]+:$',
            r'^[0-9]+\.\s+[A-Z]',
            r'^[A-Z][A-Z\s]+$'
        ]
        for pattern in heading_patterns:
            if re.match(pattern, para.text.strip()):
                return True
        return False

    def _find_table_section(self, table):
        for section, content in self.content['sections'].items():
            for item in content:
                if isinstance(item, Table):
                    if item._element is table._element:
                        return section
        return None

    def _set_metadata(self, catalog_no, lot_no):
        """Set catalog number and lot number in the document metadata and headers."""
        print(f"Setting metadata: Catalog No: {catalog_no}, Lot No: {lot_no}")
        
        # Add company name to replacements based on detected company
        company_display_names = {
            'boster': 'Boster',
            'red_dot': 'Red Dot',
            'assay_pro': 'Assay Pro',
            'icl': 'ICL',
            'elabs': 'eLabs',
            'arbor_assay': 'Arbor Assay'
        }
        
        if self.company in company_display_names:
            company_name = company_display_names[self.company]
            if all(company_name != old for old, _ in self._replacements):
                self._replacements.append((company_name, 'Innovative Research'))
        
        # Find and replace catalog number and lot number in the document
        for para in self.template.paragraphs:
            if 'Cat.' in para.text or 'Catalog' in para.text or 'Product #' in para.text:
                # Mark paragraph for catalog number
                for run in para.runs:
                    if any(x in run.text for x in ('Cat.', 'Catalog', 'Product #')):
                        # Either replace existing number or add the catalog number
                        run.text = re.sub(r'(?:Cat\.|Catalog|Product) *#?:? *[\w-]*', 
                                         f'Catalog #: {catalog_no}', 
                                         run.text)
            
            if 'Lot' in para.text:
                # Mark paragraph for lot number
                for run in para.runs:
                    if 'Lot' in run.text:
                        # Either replace existing number or add the lot number
                        run.text = re.sub(r'Lot *#?:? *[\w-]*', 
                                         f'Lot #: {lot_no}', 
                                         run.text)
        
        # Check headers and footers for metadata fields
        for section in self.template.sections:
            for header in section.header.paragraphs:
                if 'Cat.' in header.text or 'Catalog' in header.text:
                    for run in header.runs:
                        if any(x in run.text for x in ('Cat.', 'Catalog', 'Product #')):
                            run.text = re.sub(r'(?:Cat\.|Catalog|Product) *#?:? *[\w-]*', 
                                             f'Catalog #: {catalog_no}', 
                                             run.text)
                
                if 'Lot' in header.text:
                    for run in header.runs:
                        if 'Lot' in run.text:
                            run.text = re.sub(r'Lot *#?:? *[\w-]*', 
                                             f'Lot #: {lot_no}', 
                                             run.text)

    def _process_section(self, section_name, extraction_logic):
        section_content = None
        if section_name in self.content['sections']:
            section_content = self.content['sections'][section_name]
        
        if section_content:
            extracted_content = extraction_logic(section_content)
            #add logic to put extracted content into the correct place
            print(f"Processing: {section_name}")

    def _process_intended_use(self):
        self._process_section("Assay Principle", lambda x: x[0].text if x else "")

    def _process_background(self):
        self._process_section("Background on ", lambda x: x[0].text if x else "")

    def _process_assay_principle(self):
        self._process_section("Assay Principle", lambda x: x[-1].text if x else "")


    def _process_overview(self):
       self._process_section("Overview", lambda x: x)

    def _process_technical_details(self):
        self._process_section("Technical Details", lambda x: "".join([p.text for p in x]))

    def _process_preparations_before_assay(self):
        self._process_section("Preparations Before Assay", lambda x: "".join([p.text for p in x]))

    def _process_kit_components(self):
        self._process_section("Kit Components", lambda x: x) # Handle table specifically

    def _process_required_materials(self):
        self._process_section("Required Materials", lambda x: "".join([p.text for p in x]))

    def _process_standard_curve(self):
        self._process_section("Standard Curve Example", lambda x: "".join([p.text for p in x]))

    def _process_assay_variability(self):
        self._process_section("Intra/Inter-Assay Variability", lambda x: x) # Handle table specifically

    def _process_reproducibility(self):
        self._process_section("Reproducibility", lambda x: x) # Handle table specifically

    def _process_experiment_preparation(self):
        self._process_section("Preparation Before the Experiment", lambda x: x) # Handle table specifically

    def _process_standard_dilution(self):
        self._process_section("Dilution of ", lambda x: x)

    def _process_sample_preparation(self):
        self._process_section("Sample Preparation and Storage", lambda x: "".join([p.text for p in x]))

    def _process_sample_collection(self):
        self._process_section("Sample Collection Notes", lambda x: "".join([p.text for p in x]))

    def _process_sample_dilution(self):
        self._process_section("Sample Dilution Guideline", lambda x: "".join([p.text for p in x]))

    def _process_assay_protocol(self):
        self._process_section("Assay Protocol", lambda x: "".join([p.text for p in x]))

    def _process_protocol_notes(self):
        self._process_section("Assay Protocol Notes", lambda x: "".join([p.text for p in x]))

    def _process_data_analysis(self):
        self._process_section("Data Analysis", lambda x: "".join([p.text for p in x[-2:]]))

    # Original _add_disclaimer method was a placeholder


    def _apply_formatting(self):
        print("Applying document formatting")
        # Set font to Calibri 11 for body
        for para in self.template.paragraphs:
            para.style.font.name = self._styles['body']['font']
            para.style.font.size = Pt(self._styles['body']['size'])
            # Use a more compatible way to set line spacing
            para.paragraph_format.line_spacing = self._styles['body']['line_spacing']
        # Set headings to Calibri 12 bold blue
        for para in self.template.paragraphs:
            if para.style.name.startswith('Heading'):
                para.style.font.name = self._styles['header']['font']
                para.style.font.size = Pt(self._styles['header']['size'])
                para.style.font.bold = self._styles['header']['bold']
                para.style.font.color.rgb = self._styles['header']['color']
                para.paragraph_format.alignment = self._styles['header']['alignment']

    def _replace_text(self, old_text, new_text):
        print(f"Replacing: '{old_text}' with '{new_text}'")
        for para in self.template.paragraphs:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
        for table in self.template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)

    def _extract_supplier_content(self):
        current_section = None
        for para in self.supplier.paragraphs:
            if para.style.name.startswith('Heading'):
                current_section = para.text.strip()
                self.content['sections'][current_section] = []
            elif current_section:
                self._store_paragraph_content(current_section, para)

        for table in self.supplier.tables:
            self.content['tables'].append(self._process_table(table))

        for shape in self.supplier.inline_shapes:
            if shape.type == 3:  # Picture
                self.content['images'].append(shape)
    
    def _store_paragraph_content(self, section, paragraph):
        content = {
            'text': paragraph.text,
            'runs': [{
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic
            } for run in paragraph.runs]
        }
        self.content['sections'][section].append(content)

    def _process_table(self, table):
        processed = []
        for row in table.rows:
            processed_row = []
            for cell in row.cells:
                cell_content = []
                for para in cell.paragraphs:
                    cell_content.append({
                        'text': para.text,
                        'runs': [{
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic
                        } for run in para.runs]
                    })
                processed_row.append(cell_content)
            processed.append(processed_row)
        return processed
        
    def _apply_global_formatting(self):
        # Body formatting
        style = self.template.styles['Normal']
        font = style.font
        font.name = self._styles['body']['font']
        font.size = Pt(self._styles['body']['size'])
        font.color.rgb = self._styles['body']['color']
        self.template.styles['Normal'].paragraph_format.line_spacing = self._styles['body']['line_spacing']

        # Force single-column layout for all sections
        for section in self.template.sections:
            sectPr = section._sectPr
            cols = sectPr.xpath('./w:cols')
            if cols:
                cols = cols[0]
                cols.set(qn('w:num'), '1')  # Set to 1 column
                cols.set(qn('w:space'), '0')  # No spacing between columns
            else:
                cols = OxmlElement('w:cols')
                cols.set(qn('w:num'), '1')
                cols.set(qn('w:space'), '0')
                sectPr.append(cols)

    def _add_missing_reagent_rows(self, table):
        if self.company == 'boster':
            # Boster-specific reagent data
            missing_data = [
                ["Microplate", "Equilibrate at room temperature for 60 min", "Store at 4°C"],
                ["Samples", "Dilute 1:50 in provided buffer", "Aliquot and store at -80°C"]
            ]
        elif self.company == 'assay_pro':
            # Assay Pro-specific reagent data
            missing_data = [
                ["Microplate", "Equilibrate at room temperature for 30 min", "Store at 4°C"],
                ["Samples", "Dilute 1:25 in assay diluent", "Store at -20°C"]
            ]
        elif self.company == 'icl':
            # ICL-specific reagent data
            missing_data = [
                ["Microplate", "Equilibrate at room temperature", "Store at 4°C"],
                ["Samples", "Dilute as directed in protocol", "Store at -20°C"]
            ]
        elif self.company == 'elabs':
            # eLabs-specific reagent data
            missing_data = [
                ["Microplate", "Equilibrate at room temperature for 45 min", "Store at 4°C"],
                ["Samples", "Dilute 1:40 in sample buffer", "Store at -20°C"]
            ]
        elif self.company == 'arbor_assay':
            # Arbor Assay-specific reagent data
            missing_data = [
                ["Microplate", "Equilibrate at room temperature for 30 min", "Store at 4°C"],
                ["Samples", "Dilute according to sample type", "Aliquot and store at -80°C"]
            ]
        else:
            # Default values for other suppliers
            missing_data = [
                ["Microplate", "Equilibrate to room temperature", "Store at 4°C"],
                ["Samples", "Dilute 1:100 in assay buffer", "Store at -20°C"]
            ]
        
        for row_data in missing_data:
            new_row = table.add_row()
            for idx, text in enumerate(row_data):
                cell = new_row.cells[idx]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(text)
                run.italic = False
                
                # Add units if needed
                if 'temperature' in text.lower():
                    run.text += " (20-25°C)"
                elif 'buffer' in text.lower():
                    run.text += " (1X)"
        
        self._apply_table_borders(table)
    
    def _handle_reagent_section(self, section_heading, content, mapping):
        # Clear existing content
        self._clear_existing_content(section_heading)
        
        # Add original content without italics
        for item in content:
            new_para = self.template.add_paragraph()
            for run in item['runs']:
                new_run = new_para.add_run(self._apply_replacements(run['text']))
                new_run.bold = run['bold']
                new_run.italic = False  # Force disable italics
        
        # Find and format the table
        table = self._find_template_table(section_heading.text)
        if table:
            # Apply non-italic formatting to existing cells
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.italic = False
            
            # Add missing rows
            self._add_missing_reagent_rows(table)
            
            # Add units to existing cells
            self._add_units_to_table(table)
            
            # Apply table borders
            self._apply_table_borders(table)

    def _add_units_to_table(self, table):
        # Add units to specific items
        unit_mappings = {
            'assay buffer': '(1X)',
            'room temperature': '(20-25°C)',
            'store at 4': '(4°C)',
            'store at -20': '(-20°C)'
        }
        
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip().lower()
                
                # Add units where missing
                for keyword, unit in unit_mappings.items():
                    if keyword in cell_text and unit not in cell_text:
                        # Replace the cell text
                        new_text = cell.text.strip() + " " + unit
                        
                        # Clear cell and add non-italic text
                        cell.text = ""
                        for paragraph in cell.paragraphs:
                            run = paragraph.add_run(new_text)
                            run.italic = False
                            
    def _process_sections(self):
        for para in self.template.paragraphs:
            if para.style.name.startswith('Heading'):
                section_text = para.text.strip()
                if 'REAGENT PREPARATION' in section_text.upper():
                    mapping = self.mappings[self.company]['section_map'].get(section_text)
                    content = self._find_matching_content(mapping)
                    self._handle_reagent_section(para, content, mapping)
                mapping = self.mappings[self.company]['section_map'].get(section_text)
                if mapping:
                    content = self._find_matching_content(mapping)
                    if content:
                        # Special handling for Reagent section
                        if 'Reagent' in section_text:
                            self._handle_reagent_section(para, content, mapping)
                        else:
                            self._replace_section_content(para, content, mapping)
                        
    def _find_matching_content(self, mapping):
        if not mapping:
            return None
            
        pattern = re.compile(mapping['source'], re.IGNORECASE) if mapping.get('regex') else None
        for section in self.content['sections']:
            if pattern and pattern.search(section):
                return self.content['sections'][section]
            elif section == mapping.get('source'):
                return self.content['sections'][section]
        return None

    def _replace_section_content(self, section_heading, content, mapping):
        self._clear_existing_content(section_heading)
        self._add_new_content(section_heading, content, mapping)

        # Add missing rows specifically for Reagent Preparation
        if 'Reagent Preparation' in section_heading.text:
            table = self._find_template_table(section_heading.text)
            if table:
                self._add_missing_reagent_rows(table)
                
    def _clear_existing_content(self, section_heading):
        current_elem = section_heading._p
        while True:
            next_elem = current_elem.getnext()
            if not next_elem or next_elem.tag.endswith('pPr'):
                break
            if next_elem.tag.endswith('p'):
                try:
                    next_para_style = self.template.styles.get_style_by_id(next_elem.xpath('.//w:pStyle/@w:val')[0])
                    if next_para_style and next_para_style.name.startswith('Heading'):
                        break
                except:
                    pass
            current_elem.getparent().remove(next_elem)

    def _add_new_content(self, section_heading, content, mapping):
        target_paragraph = None
        if 'paragraph' in mapping:
            try:
                target_paragraph = content[mapping['paragraph']]
            except IndexError:
                pass

        for item in (content if not target_paragraph else [target_paragraph]):
            new_para = self.template.add_paragraph()
            for run in item['runs']:
                new_run = new_para.add_run(self._apply_replacements(run['text']))
                new_run.bold = run['bold']
                new_run.italic = run['italic']
            new_para._p.addnext(section_heading._p)

    def _apply_table_borders(self, table):
        # Apply consistent borders to all table cells
        tbl = table._tbl
        for cell in tbl.iter_cells():
            tcPr = cell.tcPr
            # Remove existing borders
            for tag in ['tcBorders', 'top', 'left', 'bottom', 'right']:
                element = tcPr.find(qn(f'w:{tag}'))
                if element is not None:
                    tcPr.remove(element)
            
            # Add new borders
            borders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{edge}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                borders.append(border)
            tcPr.append(borders)
                 
    def _find_template_table(self, section_name):
        for para in self.template.paragraphs:
            if para.text.strip().upper() == section_name.upper():
                # Find the next table after the section heading
                elem = para._element
                while elem is not None:
                    elem = elem.getnext()
                    if elem is not None and elem.tag.endswith('tbl'):
                        # Find the table by comparing elements
                        for table in self.template.tables:
                            if table._element is elem:
                                return table
                        break
        return None
        
    def _apply_replacements(self, text):
        for old, new in self._replacements:
            text = text.replace(old, new)
        return text
    
    def _remove_existing_disclaimers(self):
         # Define disclaimer components to match (case-insensitive)
        disclaimer_phrases = [
            "disclaimer",
            "in-vitro use only",
            "not suitable for human use",
            "sufficient verification and testing",
            "statements herein are offered for informational purposes"
        ]
        
        # Check all paragraphs and their runs
        for para in list(self.template.paragraphs):
            full_text = para.text.lower()
            
            # Check paragraph text and individual runs
            remove_para = any(phrase in full_text for phrase in disclaimer_phrases)
            
            # Check runs for italic/bold formatting that might contain disclaimer text
            if not remove_para:
                for run in para.runs:
                    run_text = run.text.lower()
                    if any(phrase in run_text for phrase in disclaimer_phrases):
                        remove_para = True
                        break
            
            if remove_para:
                try:
                    p = para._element
                    p.getparent().remove(p)
                except:
                    pass
    
    def _add_disclaimer(self):
        # Remove existing disclaimer sections first
        self._remove_existing_disclaimers()
        
        disclaimer_text = (
            "This information is believed to be correct but does not claim to be all-inclusive and shall be used only as a guide. "
            "The supplier of this kit shall not be held liable for any damage resulting from handling of or contact with the above product. "
            "\n"
            "This material is sold for in-vitro use only in manufacturing and research. "
            "This material is not suitable for human use. "
            "It is the responsibility of the user to undertake sufficient verification and testing to determine the suitability of each product's application. "
            "The statements herein are offered for informational purposes only and are intended to be used solely for your consideration, investigation and verification. "
        )
        
        disclaimer = self.template.add_paragraph()
        disclaimer_run = disclaimer.add_run('DISCLAIMER: ')
        disclaimer_run.bold = True
        disclaimer_run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue color
        disclaimer.add_run(disclaimer_text).italic = True
        disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _remove_unused_sections(self):
        company_mappings = self.mappings.get(self.company, self.mappings['default'])
        remove_sections = company_mappings.get('remove_sections', [])
        
        for section in remove_sections:
            for para in self.template.paragraphs:
                if para.text.strip() == section:
                    self._delete_section(para)

    def _delete_section(self, paragraph):
        current_elem = paragraph._p
        try:
            while True:
                next_elem = current_elem.getnext()
                if not next_elem or next_elem.tag.endswith('pPr'):
                    break
                if next_elem.tag.endswith('p'):
                    try:
                        next_para_style = self.template.styles.get_style_by_id(next_elem.xpath('.//w:pStyle/@w:val')[0])
                        if next_para_style and next_para_style.name.startswith('Heading'):
                            break
                    except:
                        pass
                current_elem.getparent().remove(next_elem)
            current_elem.getparent().remove(current_elem)
        except:
            pass
            
    def process(self):
        self._extract_supplier_content()
        print(f"Found sections: {list(self.content['sections'].keys())}")  # Debug
        print(f"Found tables: {len(self.content['tables'])}")  # Debug
        self._apply_global_formatting()
        self._process_sections()
        self._handle_tables()
        self._handle_images()
        self._add_disclaimer()
        self._remove_unused_sections()
        self._apply_formatting()
        for old, new in self._replacements:
            self._replace_text(old, new)
        self.template.save(f'converted_{self.supplier_filename}.docx')
        
    def _handle_tables(self):
        # Process tables from the Advanced DOCX converter version
        print("Processing tables...")
        
    def _handle_images(self):
        # Handle image insertion from the Advanced DOCX converter version
        curve_section = self._find_standard_curve_section()
        if curve_section and self.content['images']:
            self._insert_image(curve_section, self.content['images'][0])
            
    def _find_standard_curve_section(self):
        for para in self.template.paragraphs:
            if 'Standard Curve Example' in para.text:
                return para
        return None

    def _insert_image(self, position, image):
        run = position.add_run()
        try:
            run.add_picture(image.image.blob, width=Inches(4))
            # Add space after image
            position.add_paragraph()
        except Exception as e:
            print(f"Error inserting image: {str(e)}")


def process_elisa_document(outside_doc_path, template_doc_path, catalog_no, lot_no, company=None):
    """
    Process an ELISA kit document with the enhanced processor.

    Args:
        outside_doc_path (str): Path to the outside document
        template_doc_path (str): Path to the template document
        catalog_no (str): Catalog number
        lot_no (str): Lot number
        company (str): Company identifier

    Returns:
        str: Path to the generated document
    """
    processor = ELISAProcessor(template_doc_path, outside_doc_path)
    
    # Extract the supplier content
    processor._extract_supplier_content()
    print(f"Found sections: {list(processor.content['sections'].keys())}")  # Debug
    print(f"Found tables: {len(processor.content['tables'])}")  # Debug
    
    # Apply global formatting
    processor._apply_global_formatting()
    
    # Set metadata (catalog number and lot number)
    processor._set_metadata(catalog_no, lot_no)
    
    # Process sections
    processor._process_sections()
    
    # Handle tables
    processor._handle_tables()
    
    # Handle images
    processor._handle_images()
    
    # Add disclaimer
    processor._add_disclaimer()
    
    # Remove unused sections
    processor._remove_unused_sections()
    
    # Apply formatting
    processor._apply_formatting()
    
    # Replace text
    for old, new in processor._replacements:
        processor._replace_text(old, new)
    
    # Save the document
    output_path = f'converted_{processor.supplier_filename}.docx'
    processor.template.save(output_path)
    
    return output_path

if __name__ == "__main__":
    print("ELISA Document Converter")
    print("------------------------")

    template_path = input("Enter path to template document: ")
    supplier_path = input("Enter path to supplier document: ")
    catalog_no = input("Enter catalog number: ")
    lot_no = input("Enter lot number: ")
    
    output_path = process_elisa_document(supplier_path, template_path, catalog_no, lot_no)
    print(f"Document processing completed successfully!")
    print(f"Output saved to: {output_path}")