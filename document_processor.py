import logging
from docx import Document

logger = logging.getLogger(__name__)

def process_document(input_path, output_path, find_text, replace_text):
    """
    Process a DOCX document by replacing text.
    
    Args:
        input_path (str): Path to the input DOCX file
        output_path (str): Path to save the processed DOCX file
        find_text (str): Text to find
        replace_text (str): Text to replace with
        
    Returns:
        int: Number of replacements made
    """
    logger.debug(f"Processing document: {input_path}")
    logger.debug(f"Find text: '{find_text}', Replace with: '{replace_text}'")
    
    try:
        # Load the document
        doc = Document(input_path)
        
        # Counter for replacements
        replacement_count = 0
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                for run in paragraph.runs:
                    if find_text in run.text:
                        run.text = run.text.replace(find_text, replace_text)
                        replacement_count += run.text.count(replace_text)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if find_text in paragraph.text:
                            for run in paragraph.runs:
                                if find_text in run.text:
                                    run.text = run.text.replace(find_text, replace_text)
                                    replacement_count += run.text.count(replace_text)
        
        # Process headers and footers
        for section in doc.sections:
            # Process headers
            for header in section.header.paragraphs:
                if find_text in header.text:
                    for run in header.runs:
                        if find_text in run.text:
                            run.text = run.text.replace(find_text, replace_text)
                            replacement_count += run.text.count(replace_text)
            
            # Process footers
            for footer in section.footer.paragraphs:
                if find_text in footer.text:
                    for run in footer.runs:
                        if find_text in run.text:
                            run.text = run.text.replace(find_text, replace_text)
                            replacement_count += run.text.count(replace_text)
        
        # Save the modified document
        doc.save(output_path)
        
        logger.debug(f"Made {replacement_count} replacements")
        return replacement_count
    
    except Exception as e:
        logger.exception(f"Error processing document: {str(e)}")
        raise
