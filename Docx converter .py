import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement, qn

class ELISAProcessor:
    def __init__(self, template_path, supplier_path):
        self.template = Document(template_path)
        self.supplier = Document(supplier_path)
        self.supplier_filename = os.path.splitext(os.path.basename(supplier_path))[0]  # get 'supplier1' from 'supplier1.docx'
        self.company = self._detect_company(supplier_path)
        self.content = {'sections': {}, 'tables': [], 'images': []}
        self.mappings = self._load_mappings()
        self._styles = {
            'header': {'font': 'Calibri', 'size': 24, 'bold': True, 
                      'color': RGBColor(0x00, 0x00, 0xFF), 'alignment': WD_ALIGN_PARAGRAPH.RIGHT},
            'body': {'font': 'Calibri', 'size': 11, 
                     'line_spacing': 1.15, 'color': RGBColor(0x00, 0x00, 0x00)},
            'disclaimer': {'italic': True}
        }
        self._replacements = [('Boster', 'Innovative Research')]

    def _detect_company(self, path):
        filename = os.path.basename(path)
        if filename.startswith('EK'): return 'boster'
        if filename.startswith('RDR'): return 'red_dot'
        return list(self._load_mappings().keys())[0]

    def _load_mappings(self):
        return {
            'boster': {
                'section_map': {
                    'INTENDED USE': {'source': r'Assay Principle', 'regex': True, 'paragraph': 0},
                    'BACKGROUND': {'source': r'Background on', 'regex': True},
                    'ASSAY PRINCIPLE': {'source': r'Assay Principle', 'regex': True, 'paragraph': -1},
                    'OVERVIEW': {'source': 'Overview', 'table_filter': ['Product Name', 'Description']},
                    'TECHNICAL DETAILS': {'source': 'Technical Details'},
                    'STANDARD CURVE EXAMPLE': {'source': r'Standard Curve Example$', 'regex': True},
                    'INTRA/INTER-ASSAY VARIABILITY': {
                    'source': r'(Intra|Inter).*Variability',
                    'regex': True,
                    'table_headers': ['Intra-Assay', 'Inter-Assay']
                },
                'REPRODUCIBILITY': {
                    'source': r'Reproducibility',
                    'table_headers': ['Reproducibility']
                },
                'REAGENT PREPARATION AND STORAGE': {
                    'source': 'Reagent Preparation',
                    'special_handling': True,
                    'table_headers': ['Component', 'Preparation', 'Storage']
                }
                    # Add other sections as needed
                },
                'remove_sections': ['Dilution of Standard', 'Sample Collection Notes', 'Sample Activation']
            },
            'red_dot': {
                'section_map': {
                    'INTENDED USE': {'source': r'Assay Principle', 'regex': True, 'paragraph': 0},
                    'BACKGROUND': {'source': r'Background on', 'regex': True},
                    'ASSAY PRINCIPLE': {'source': r'Assay Principle', 'regex': True, 'paragraph': -1},
                    'OVERVIEW': {'source': 'Overview', 'table_filter': ['Product Name', 'Description']},
                    'TECHNICAL DETAILS': {'source': 'Technical Details'},
                    'STANDARD CURVE EXAMPLE': {'source': r'Standard Curve Example$', 'regex': True},
                    'INTRA/INTER-ASSAY VARIABILITY': {
                    'source': r'(Intra|Inter).*Variability',
                    'regex': True,
                    'table_headers': ['Intra-Assay', 'Inter-Assay']
                },
                'REPRODUCIBILITY': {
                    'source': r'Reproducibility',
                    'table_headers': ['Reproducibility']
                }
                    # Add other sections as needed
                },
                'remove_sections': ['Dilution of Standard', 'Sample Collection Notes', 'Sample Activation']
            },
            'default': {'section_map': {}, 'remove_sections': []}
        }

    def _extract_supplier_content(self):
        current_section = None
        for para in self.supplier.paragraphs:
            if para.style.name.startswith('Heading'):
                current_section = para.text.strip()
                self.content['sections'][current_section] = []
            elif current_section:
                self._store_paragraph_content(current_section, para)

        for table in self.supplier.tables:
            self.content['tables'].append(self._process_table(table))

        for shape in self.supplier.inline_shapes:
            if shape.type == 3:
                self.content['images'].append(shape)

    def _store_paragraph_content(self, section, paragraph):
        content = {
            'text': paragraph.text,
            'runs': [{
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic
            } for run in paragraph.runs]
        }
        self.content['sections'][section].append(content)

    def _process_table(self, table):
        processed = []
        for row in table.rows:
            processed_row = []
            for cell in row.cells:
                cell_content = []
                for para in cell.paragraphs:
                    cell_content.append({
                        'text': para.text,
                        'runs': [{
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic
                        } for run in para.runs]
                    })
                processed_row.append(cell_content)
            processed.append(processed_row)
        return processed

    def _apply_global_formatting(self):
        
        # Body formatting
        style = self.template.styles['Normal']
        font = style.font
        font.name = self._styles['body']['font']
        font.size = Pt(self._styles['body']['size'])
        font.color.rgb = self._styles['body']['color']
        self.template.styles['Normal'].paragraph_format.line_spacing = self._styles['body']['line_spacing']

        # Force single-column layout for all sections
        for section in self.template.sections:
            sectPr = section._sectPr
            cols = sectPr.xpath('./w:cols')
            if cols:
                cols = cols[0]
                cols.set(qn('w:num'), '1')  # Set to 1 column
                cols.set(qn('w:space'), '0')  # No spacing between columns
            else:
                cols = OxmlElement('w:cols')
                cols.set(qn('w:num'), '1')
                cols.set(qn('w:space'), '0')
                sectPr.append(cols)

    def _add_missing_reagent_rows(self, table):
        if self.company == 'boster':
            # Boster-specific reagent data
            missing_data = [
                ["Microplate", "Equilibrate at room temperature for 60 min", "Store at 4°C"],
                ["Samples", "Dilute 1:50 in provided buffer", "Aliquot and store at -80°C"]
            ]
        else:
            # Default values for other suppliers
            missing_data = [
                ["Microplate", "Equilibrate to room temperature", "Store at 4°C"],
                ["Samples", "Dilute 1:100 in assay buffer", "Store at -20°C"]
            ]
        
        for row_data in missing_data:
            new_row = table.add_row()
            for idx, text in enumerate(row_data):
                cell = new_row.cells[idx]
                cell.text = ""
                para = cell.paragraphs[0]
                run = para.add_run(text)
                run.italic = False
                
                # Add units if needed
                if 'temperature' in text.lower():
                    run.text += " (20-25°C)"
                elif 'buffer' in text.lower():
                    run.text += " (1X)"
        
        self._apply_table_borders(table)

    
    def _handle_reagent_section(self, section_heading, content, mapping):
        # Clear existing content
        self._clear_existing_content(section_heading)
        
        # Add original content without italics
        for item in content:
            new_para = self.template.add_paragraph()
            for run in item['runs']:
                new_run = new_para.add_run(self._apply_replacements(run['text']))
                new_run.bold = run['bold']
                new_run.italic = False  # Force disable italics
        
        # Find and format the table
        table = self._find_template_table(section_heading.text)
        if table:
            # Apply non-italic formatting to existing cells
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.italic = False
            
            # Add missing rows
            self._add_missing_reagent_rows(table)
            
            # Add units to existing cells
            self._add_units_to_table(table)
            
            # Apply table borders
            self._apply_table_borders(table)

    def _add_units_to_table(self, table):
        # Add units to specific items
        unit_mappings = {
            'assay buffer': '(1X)',
            'room temperature': '(20-25°C)',
            'store at 4': '(4°C)',
            'store at -20': '(-20°C)'
        }
        
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip().lower()
                
                # Add units where missing
                for keyword, unit in unit_mappings.items():
                    if keyword in cell_text and unit not in cell_text:
                        # Replace the cell text
                        new_text = cell.text.strip() + " " + unit
                        
                        # Clear cell and add non-italic text
                        cell.text = ""
                        for paragraph in cell.paragraphs:
                            run = paragraph.add_run(new_text)
                            run.italic = False

    def _apply_guaranteed_borders(self, table):
        """Force borders on all table cells and edges"""
        tbl = table._tbl
        
        # Table-level borders
        tbl_pr = tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        
        # Replace existing borders
        existing = tbl_pr.xpath('.//w:tblBorders')
        for element in existing:
            element.getparent().remove(element)
        tbl_pr.append(borders)

        # Cell-level borders
        for cell in tbl.iter_cells():
            tc_pr = cell.tcPr
            cell_borders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{edge}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                cell_borders.append(border)
            
            # Remove existing cell borders
            existing_cell = tc_pr.xpath('.//w:tcBorders')
            for element in existing_cell:
                element.getparent().remove(element)
            tc_pr.append(cell_borders)

    

    def _process_sections(self):
        for para in self.template.paragraphs:
            if para.style.name.startswith('Heading'):
                section_text = para.text.strip()
                if 'REAGENT PREPARATION' in section_text.upper():
                    mapping = self.mappings[self.company]['section_map'].get(section_text)
                    content = self._find_matching_content(mapping)
                    self._handle_reagent_section(para, content, mapping)
                mapping = self.mappings[self.company]['section_map'].get(section_text)
                if mapping:
                    content = self._find_matching_content(mapping)
                    if content:
                        # Special handling for Reagent section
                        if 'Reagent' in section_text:
                            self._handle_reagent_section(para, content, mapping)
                        else:
                            self._replace_section_content(para, content, mapping)
                        
                        

    def _find_matching_content(self, mapping):
        pattern = re.compile(mapping['source'], re.IGNORECASE) if mapping.get('regex') else None
        for section in self.content['sections']:
            if pattern and pattern.search(section):
                return self.content['sections'][section]
            elif section == mapping.get('source'):
                return self.content['sections'][section]
        return None

    def _replace_section_content(self, section_heading, content, mapping):
        self._clear_existing_content(section_heading)
        self._add_new_content(section_heading, content, mapping)

        # Add missing rows specifically for Reagent Preparation
        if 'Reagent Preparation' in section_heading.text:
            self._add_missing_reagent_rows(section_heading)
    
    def _add_missing_reagent_rows(self, table):
        missing_data = [
            ["Microplate", "Equilibrate to room temperature", "Store at 4°C"],
            ["Samples", "Dilute 1:100 in assay buffer", "Store at -20°C"]
        ]
        
        for row_data in missing_data:
            new_row = table.add_row()
            for idx, text in enumerate(row_data):
                cell = new_row.cells[idx]
                cell.text = ''
                para = cell.paragraphs[0]
                run = para.add_run(text)
                run.italic = False  # Explicitly remove italics
                # Add units dynamically
                if 'temperature' in text.lower():
                    run.text += " (20-25°C)"
                elif 'buffer' in text.lower():
                    run.text += " (1X)"
        
        self._apply_table_borders(table)  # Ensure borders are applied
    
    def _format_reagent_cell(self, cell, text):
        cell.text = ''
        para = cell.paragraphs[0]
        
        # Add unit conversion if needed
        if 'buffer' in text.lower():
            text += " (1X)"
        elif 'temperature' in text.lower():
            text += " (20-25°C)"
        
        run = para.add_run(text)
        run.bold = False
        run.italic = False  # Ensure not italicized
        
        # Set table borders
        tbl = cell._tc.get_or_add_tbl()
        tbl.tblPr.xpath('./w:tblBorders')[0].remove()
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), 'auto')
            borders.append(edge_element)
        tbl.tblPr.append(borders)

    def _clear_existing_content(self, section_heading):
        current_elem = section_heading._p
        while True:
            next_elem = current_elem.getnext()
            if not next_elem or next_elem.tag.endswith('pPr'):
                break
            if next_elem.tag.endswith('p'):
                next_para_style = self.template.styles.get_style_by_id(next_elem.xpath('.//w:pStyle/@w:val')[0])
                if next_para_style and next_para_style.name.startswith('Heading'):
                    break
            current_elem.getparent().remove(next_elem)

    def _add_new_content(self, section_heading, content, mapping):
        target_paragraph = None
        if 'paragraph' in mapping:
            try:
                target_paragraph = content[mapping['paragraph']]
            except IndexError:
                pass

        for item in (content if not target_paragraph else [target_paragraph]):
            new_para = self.template.add_paragraph()
            for run in item['runs']:
                new_run = new_para.add_run(self._apply_replacements(run['text']))
                new_run.bold = run['bold']
                new_run.italic = run['italic']
            new_para._p.addnext(section_heading._p)

    def _apply_complete_table_borders(self, table):
        # Set table-wide borders
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        
        # Clear existing borders
        existing_borders = tbl_pr.xpath('.//w:tblBorders')
        for border in existing_borders:
            border.getparent().remove(border)
        
        tbl_pr.append(borders)

        # Set cell-level borders
        for cell in tbl.iter_cells():
            tc_pr = cell.tcPr
            cell_borders = OxmlElement('w:tcBorders')
            
            for edge in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{edge}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                cell_borders.append(border)
            
            # Clear existing cell borders
            existing_cell_borders = tc_pr.xpath('.//w:tcBorders')
            for border in existing_cell_borders:
                border.getparent().remove(border)
            
            tc_pr.append(cell_borders)

    def _apply_table_borders(self, table):
        # Apply consistent borders to all table cells
        tbl = table._tbl
        for cell in tbl.iter_cells():
            tcPr = cell.tcPr
            # Remove existing borders
            for tag in ['tcBorders', 'top', 'left', 'bottom', 'right']:
                element = tcPr.find(qn(f'w:{tag}'))
                if element is not None:
                    tcPr.remove(element)
            
            # Add new borders
            borders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{edge}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                borders.append(border)
            tcPr.append(borders)
     
    def _find_template_table(self, section_name):
        for para in self.template.paragraphs:
            if para.text.strip().upper() == section_name.upper():
                # Find the next table after the section heading
                elem = para._element
                while elem is not None:
                    elem = elem.getnext()
                    if elem is not None and elem.tag.endswith('tbl'):
                        # Find the table by comparing elements
                        for table in self.template.tables:
                            if table._element is elem:
                                return table
                        break
        return None
    
    def _process_mapped_tables(self):
        for section_name, mapping in self.mappings[self.company]['section_map'].items():
            if 'table_headers' in mapping:
                template_table = self._find_template_table(section_name)
                if template_table:
                    supplier_table = self._find_supplier_table_by_mapping(mapping)
                    if supplier_table:
                        self._replace_table_content(template_table, supplier_table)

    def _find_supplier_table_by_mapping(self, mapping):
        target_headers = [h.lower() for h in mapping.get('table_headers', [])]
        
        for table_idx, table in enumerate(self.content['tables']):
            if not table or not table[0]:
                continue
                
            # Check first row cells for target headers
            header_cells = []
            for cell in table[0]:
                cell_text = ' '.join([p['text'].strip() for p in cell]).lower()
                header_cells.append(cell_text)
                print(f"Table {table_idx} header cell: {cell_text}")  # Debug log

            # Check for any target header in any cell
            match_score = sum(
                1 for th in target_headers
                for hc in header_cells
                if th in hc
            )
            
            if match_score >= len(target_headers):
                print(f"Matched table {table_idx} for {mapping['table_headers']}")  # Debug log
                return table
        
        print(f"No table found for {mapping['table_headers']}")  # Debug log
        return None

    def _handle_tables(self):
        # Process special tables first
        self._process_mapped_tables()
        
        # Then handle other tables
        processed_tables = set()
        for template_table in self.template.tables:
            try:
                if template_table in processed_tables:
                    continue
                    
                header_text = ' '.join([cell.text for cell in template_table.rows[0].cells]).lower()
                print(f"Processing table with header: {header_text}")  # Debug log
                
                supplier_table = self._find_matching_table(header_text)
                if supplier_table:
                    self._replace_table_content(template_table, supplier_table)
                    processed_tables.add(template_table)
            except Exception as e:
                print(f"Skipping table processing: {str(e)}")
                continue

    def _find_matching_table(self, header_text):
        for table in self.content['tables']:
            if table and table[0]:
                # Flatten first row cells into single string
                supplier_header = ' '.join([' '.join(p['text'] for cell in table[0] for p in cell)])
                if header_text in supplier_header.lower():
                    return table
        return None

    def _process_volume_content(self, cell, content):
        cell.text = ''
        volume_pattern = r'(\d+\s*?(?:μL|mL|ml|L|×\s*\d+))'
        found_volume = False
        
        for para in content:
            full_text = ' '.join([run['text'] for run in para['runs']])
            matches = re.findall(volume_pattern, full_text)
            
            if matches:
                # Take first volume match and format it
                volume = matches[0]
                new_para = cell.add_paragraph()
                run = new_para.add_run(volume)
                found_volume = True
                break
        
        if not found_volume:
            # Fallback to original content if no volume found
            self._fill_cell_content(cell, content)
            s
    def _replace_table_content(self, template_table, supplier_table):
        # Get template headers
        template_headers = [cell.text.strip().lower() for cell in template_table.rows[0].cells]
        
        # Check if this is the REAGENTS table
        is_reagents_table = False
        prev_para = template_table._element.getprevious()
        while prev_para is not None:
            if prev_para.tag.endswith('p'):
                para_text = ''.join(node.text for node in prev_para.iter() if node.text)
                if 'REAGENTS' in para_text.upper():
                    is_reagents_table = True
                    break
            prev_para = prev_para.getprevious()
        # Get supplier headers from first row
        supplier_headers = []
        if supplier_table and supplier_table[0]:
            supplier_headers = [
                ' '.join(p['text'].strip() for p in cell).lower()
                for cell in supplier_table[0]
            ]

        # Create column mapping with flexible matching
        column_map = {}
        
        for t_idx, t_header in enumerate(template_headers):
            # Special handling for "quantity" column that should show volumes
            if is_reagents_table and t_header == 'quantity':
                # Look for volume column in supplier headers
                volume_columns = [i for i, h in enumerate(supplier_headers) 
                                if 'volume' in h or 'μl' in h or 'ml' in h]
                if volume_columns:
                    column_map[t_idx] = volume_columns[0]
                    continue
            # Try exact match first
            if t_header in supplier_headers:
                column_map[t_idx] = supplier_headers.index(t_header)
                continue
                
            # Try partial match
            for s_idx, s_header in enumerate(supplier_headers):
                if t_header in s_header or s_header in t_header:
                    column_map[t_idx] = s_idx
                    break
            else:
                # Fallback to position-based mapping
                if t_idx < len(supplier_headers):
                    column_map[t_idx] = t_idx

        print(f"Column mapping: {column_map}")  # Debug log

        # Clear existing data rows
        for row in list(template_table.rows)[1:]:
            row._element.getparent().remove(row._element)

        # Add new rows with error handling
        for s_row in supplier_table[1:]:
            try:
                new_row = template_table.add_row()
                for t_col, s_col in column_map.items():
                    if s_col < len(s_row):
                        cell_content = s_row[s_col]
                        if is_reagents_table and template_headers[t_col] == 'quantity':
                            self._process_volume_content(new_row.cells[t_col], cell_content)
                        else:
                            self._fill_cell_content(new_row.cells[t_col], s_row[s_col])
                    else:
                        print(f"Warning: Column {s_col} out of range in supplier row")
            except Exception as e:
                print(f"Error adding row: {str(e)}")
                continue

    def _fill_cell_content(self, cell, content):
        cell.text = ''
        for para in content:
            new_para = cell.add_paragraph()
            for run in para['runs']:
                new_run = new_para.add_run(run['text'])
                new_run.bold = run.get('bold', False)
                new_run.italic = run.get('italic', False)

    def _get_cell_text(self, cell_content):
        return ' '.join([p['text'] for p in cell_content])

    def _handle_images(self):
        curve_section = self._find_standard_curve_section()
        if curve_section and self.content['images']:
            self._insert_image(curve_section, self.content['images'][0])

    def _find_standard_curve_section(self):
        for para in self.template.paragraphs:
            if 'Standard Curve Example' in para.text:
                return para
        return None

    def _insert_image(self, position, image):
        run = position.add_run()
        try:
            run.add_picture(image.image.blob, width=Inches(4))
            # Add space after image
            position.add_paragraph()
        except Exception as e:
            print(f"Error inserting image: {str(e)}")

    def _apply_replacements(self, text):
        for old, new in self._replacements:
            text = text.replace(old, new)
        return text
    
    def _remove_existing_disclaimers(self):
         # Define disclaimer components to match (case-insensitive)
        disclaimer_phrases = [
            "disclaimer",
            "in-vitro use only",
            "not suitable for human use",
            "sufficient verification and testing",
            "statements herein are offered for informational purposes"
        ]
        
        # Check all paragraphs and their runs
        for para in list(self.template.paragraphs):
            full_text = para.text.lower()
            
            # Check paragraph text and individual runs
            remove_para = any(phrase in full_text for phrase in disclaimer_phrases)
            
            # Check runs for italic/bold formatting that might contain disclaimer text
            if not remove_para:
                for run in para.runs:
                    run_text = run.text.lower()
                    if any(phrase in run_text for phrase in disclaimer_phrases):
                        remove_para = True
                        break
            
            if remove_para:
                p = para._element
                p.getparent().remove(p)
    
    def _add_disclaimer(self):
        # Remove existing disclaimer sections first
        self._remove_existing_disclaimers()
        
        disclaimer_text = (
            "This information is believed to be correct but does not claim to be all-inclusive and shall be used only as a guide."
            "The supplier of this kit shall not be held liable for any damage resulting from handling of or contact with the above product."
            "\n"
            "This material is sold for in-vitro use only in manufacturing and research."
            "This material is not suitable for human use."
            "It is the responsibility of the user to undertake sufficient verification and testing to determine the suitability of each product’s application."
            "The statements herein are offered for informational purposes only and are intended to be used solely for your consideration, investigation and verification. "
        )
        
        disclaimer = self.template.add_paragraph()
        disclaimer_run = disclaimer.add_run('DISCLAIMER: ')
        disclaimer_run.bold = True
        disclaimer_run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue color
        disclaimer.add_run(disclaimer_text).italic = True
        disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _remove_unused_sections(self):
        company_mappings = self.mappings.get(self.company, self.mappings['default'])
        remove_sections = company_mappings.get('remove_sections', [])
        
        for section in remove_sections:
            for para in self.template.paragraphs:
                if para.text.strip() == section:
                    self._delete_section(para)

    def _delete_section(self, paragraph):
        current_elem = paragraph._p
        while True:
            next_elem = current_elem.getnext()
            if not next_elem or next_elem.tag.endswith('pPr'):
                break
            if next_elem.tag.endswith('p'):
                next_para_style = self.template.styles.get_style_by_id(next_elem.xpath('.//w:pStyle/@w:val')[0])
                if next_para_style and next_para_style.name.startswith('Heading'):
                    break
            current_elem.getparent().remove(next_elem)
        current_elem.getparent().remove(current_elem)

    def process(self):
        self._extract_supplier_content()
        print(f"Found sections: {list(self.content['sections'].keys())}")  # Debug
        print(f"Found tables: {len(self.content['tables'])}")  # Debug
        self._apply_global_formatting()
        self._process_sections()
        self._handle_tables()
        self._handle_images()
        self._add_disclaimer()
        self._remove_unused_sections()
        self.template.save(f'converted_{self.supplier_filename}.docx')

if __name__ == "__main__":
    template = input("Enter Template file path: \n")
    supplier = input("Enter Supplier file path: \n")
    processor = ELISAProcessor(template, supplier)
    processor.process()
    print("Document processing completed successfully!")